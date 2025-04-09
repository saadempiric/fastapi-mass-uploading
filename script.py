import datetime
import os
import re
import requests
import json
from docx import Document
from bs4 import BeautifulSoup
from html import escape
from dotenv import load_dotenv

load_dotenv()

strapi_endpoint = os.getenv("STRAPI_ENDPOINT")
next_endpoint = os.getenv("NEXT_ENDPOINT")
api_token = os.getenv("STRAPI_API_TOKEN")

class StrapiDocUploader:
    def __init__(self, api_url, api_token = api_token):
        """
        Initialize the uploader with Strapi API credentials
        
        Args:
            api_url (str): Base URL for the Strapi API (e.g., 'http://localhost:1337/api')
            api_token (str): JWT token for authentication
        """
        self.api_url = api_url
        self.headers = {
            'Authorization': f'Bearer {api_token}',
            'Content-Type': 'application/json'
        }
    
    def _calculate_reading_time(self, content):
        """
        Calculate estimated reading time based on word count
        
        Args:
            content (str): HTML content of the blog
            
        Returns:
            int: Estimated reading time in minutes
        """
        # Remove HTML tags to count actual words
        soup = BeautifulSoup(content, 'html.parser')
        text = soup.get_text()
        
        # Average reading speed is about 200-250 words per minute
        words = len(re.findall(r'\w+', text))
        reading_time = max(1, round(words / 225))  # Minimum 1 minute, round to nearest minute
        
        return reading_time
    
    def extract_footer_text(self, doc):
        """
        Extract text from document footer, including all text elements
        
        Args:
            doc: Document object from python-docx
            
        Returns:
            str: Extracted footer text or empty string if not found
        """
        footer_text = ""
        try:
            # Loop through all sections to get their footers
            for section in doc.sections:
                # Check all footer types
                for footer_type in ['first_page_footer', 'footer', 'even_page_footer']:
                    try:
                        footer = getattr(section, footer_type)
                        if footer:
                            # Extract text from all paragraphs
                            for paragraph in footer.paragraphs:
                                if paragraph.text.strip():
                                    footer_text += paragraph.text.strip() + " "
                            
                            # Extract text from tables if present
                            for table in footer.tables:
                                for row in table.rows:
                                    for cell in row.cells:
                                        for paragraph in cell.paragraphs:
                                            if paragraph.text.strip():
                                                footer_text += paragraph.text.strip() + " "
                    except AttributeError:
                        # Skip if this footer type doesn't exist
                        continue
            
            footer_text = footer_text.strip()
            print(f"Extracted footer text: '{footer_text}'")
            return footer_text
        except Exception as e:
            print(f"Error extracting footer: {str(e)}")
            return ""
    
    def parse_doc_file(self, file_path):
        """
        Parse a DOCX file according to the required format
        
        Args:
            file_path (str): Path to the DOCX file
            
        Returns:
            dict: Parsed blog data with title, tagline, keywords, and content
        """
        try:
            doc = Document(file_path)
            
            # Initialize variables
            blog_data = {
                'title': '',
                'tagline': '',
                'keywords': [],
                'content': '',
                'modified_date': None,
                'reading_time': 0,
                'label': '' 
            }

             # Get file's last modified date
            modified_timestamp = os.path.getmtime(file_path)
            blog_data['modified_date'] = datetime.datetime.fromtimestamp(modified_timestamp).isoformat()

             # Extract footer text for the label field
            blog_data['label'] = self.extract_footer_text(doc)
            
            # Debug: Print all paragraphs with their styles to help troubleshoot
            print(f"Document contains {len(doc.paragraphs)} paragraphs")
            
            # Find and extract keywords
            keywords_found = False
            for i, para in enumerate(doc.paragraphs):
                # Print paragraph style and first few characters for debugging
                style_name = para.style.name if para.style else "No Style"
                preview = para.text[:50] + "..." if len(para.text) > 50 else para.text
                print(f"Paragraph {i}: Style '{style_name}' - '{preview}'")
                
                # Look for keywords in any paragraph containing the keywords identifier
                if "Content Keywords:" in para.text:
                    keywords_text = para.text.replace("Content Keywords:", "").strip()
                    # Split by commas and strip whitespace
                    blog_data['keywords'] = [k.strip() for k in keywords_text.split(',')]
                    keywords_found = True
                    print(f"Found keywords at paragraph {i}: {blog_data['keywords']}")
            
            if not keywords_found:
                print("WARNING: Keywords not found in document")
            
            # Process paragraphs for title, tagline, and content
            title_found = False
            tagline_found = False
            content_paragraphs = []
            content_started = False
            
            for i, para in enumerate(doc.paragraphs):
                para_text = para.text.strip()
                if not para_text:  # Skip empty paragraphs
                    continue
                
                # Skip document type indicator (e.g., 'Blog' in Book Title style)
                if i == 0 and "Blog" in para_text:
                    print(f"Skipping document type indicator: '{para_text}'")
                    continue
                
                # Skip keywords paragraph as we've already processed it
                if "Content Keywords:" in para_text:
                    continue
                
                # First Heading 1 after keywords is the title
                if not title_found and not tagline_found:
                    if para.style.name == 'Heading 1' or 'heading 1' in para.style.name.lower():
                        blog_data['title'] = para_text
                        title_found = True
                        print(f"Found title at paragraph {i}: '{para_text}'")
                        continue
                    
                # First Heading 2 after title is the tagline
                elif title_found and not tagline_found:
                    if para.style.name == 'Heading 2' or 'heading 2' in para.style.name.lower():
                        blog_data['tagline'] = para_text
                        tagline_found = True
                        content_started = True  # Content starts after tagline
                        print(f"Found tagline at paragraph {i}: '{para_text}'")
                        continue
                    # If no specific Heading 2 style is found, the next paragraph after title might be tagline
                    else:
                        blog_data['tagline'] = para_text
                        tagline_found = True
                        content_started = True
                        print(f"Assumed tagline at paragraph {i}: '{para_text}'")
                        continue
                
                # Everything after tagline is content
                elif content_started:
                    # Process paragraph based on its style for HTML conversion
                    if para.style.name == 'Heading 3' or 'heading 3' in para.style.name.lower():
                        content_paragraphs.append(f"<h3>{escape(para_text)}</h3>")
                    elif para.style.name == 'Heading 4' or 'heading 4' in para.style.name.lower():
                        content_paragraphs.append(f"<h4>{escape(para_text)}</h4>")
                    elif 'list' in para.style.name.lower():
                        content_paragraphs.append(f"<li>{self._process_inline_formatting(para)}</li>")
                    else:
                        # Regular paragraph
                        content_paragraphs.append(f"<p>{self._process_inline_formatting(para)}</p>")
            
            # If we didn't find a tagline but have title, use the next paragraph after title as content
            if title_found and not tagline_found and not content_paragraphs:
                print("WARNING: No tagline found, content may be incorrect")
            
            # Special handling if we don't find proper title/tagline
            if not title_found:
                print("WARNING: No title found in expected format")
                # Try to use the first non-empty paragraph as title
                for para in doc.paragraphs:
                    if para.text.strip() and "Content Keywords:" not in para.text and "Blog" not in para.text:
                        blog_data['title'] = para.text.strip()
                        print(f"Using '{blog_data['title']}' as fallback title")
                        break
            
            # Structure HTML content
            html_content = '\n'.join(content_paragraphs)
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Wrap consecutive li elements in ul tags
            current_list = None
            for li in soup.find_all('li'):
                if current_list is None or li.previous_sibling != current_list:
                    current_list = soup.new_tag('ul')
                    li.wrap(current_list)
                else:
                    current_list.append(li)
            
            blog_data['content'] = str(soup)
            
            # Debug output of parsed data
            print("\nParsed Blog Data:")
            print(f"Title: {blog_data['title']}")
            print(f"Tagline: {blog_data['tagline']}")
            print(f"Keywords: {blog_data['keywords']}")
            print(f"Label (from footer): {blog_data['label']}")
            print(f"Content length: {len(blog_data['content'])} characters")

            blog_data['reading_time'] = self._calculate_reading_time(blog_data['content'])
            
            return blog_data
            
        except Exception as e:
            print(f"Error parsing document {file_path}: {str(e)}")
            import traceback
            traceback.print_exc()
            return None
    
    def _process_inline_formatting(self, paragraph):
        """Process inline formatting (bold, italic, etc.) in a paragraph"""
        result = ""
        for run in paragraph.runs:
            text = escape(run.text)
            if run.bold and run.italic:
                text = f"<strong><em>{text}</em></strong>"
            elif run.bold:
                text = f"<strong>{text}</strong>"
            elif run.italic:
                text = f"<em>{text}</em>"
            result += text
        return result
    
    def upload_to_strapi(self, blog_data):
        """
        Upload parsed blog data to Strapi
        
        Args:
            blog_data (dict): Blog data with title, tagline, keywords, and content
            
        Returns:
            dict: Response from Strapi API
        """
        # Validate data before uploading
        if not blog_data['title']:
            print("ERROR: Cannot upload blog with empty title")
            return None
        
        # Prepare payload for Strapi
        payload = {
            "data": {
                "title": blog_data['title'],
                "tagline": blog_data['tagline'],
                "keywords": ','.join(blog_data['keywords']),
                "content": blog_data['content'],
                "modified_date": blog_data['modified_date'],
                "reading_time": blog_data['reading_time'],
                "label": blog_data['label']
            }
        }
        
        print(f"Uploading blog: '{blog_data['title']}'")
        
        try:
            response = requests.post(f"{self.api_url}/blogs", 
                                    headers=self.headers, 
                                    data=json.dumps(payload))
            
            if response.status_code in [200, 201]:
                print(f"Successfully uploaded blog: {blog_data['title']}")
                print(f"Modified Date: {blog_data['modified_date']}")
                print(f"Reading Time: {blog_data['reading_time']} minutes")
                print(f"Label: {blog_data['label']}")
                return response.json()
            else:
                print(f"Failed to upload blog: {response.status_code} - {response.text}")
                return None
                
        except Exception as e:
            print(f"Error uploading to Strapi: {str(e)}")
            return None
    
    def process_directory(self, directory_path):
        """
        Process all DOCX files in a directory
        
        Args:
            directory_path (str): Path to directory containing DOCX files
            
        Returns:
            list: Results of all upload attempts
        """
        results = []
        
        for filename in os.listdir(directory_path):
            if filename.lower().endswith('.docx'):
                file_path = os.path.join(directory_path, filename)
                print(f"\n========= Processing {filename} =========")
                
                blog_data = self.parse_doc_file(file_path)
                if blog_data:
                    result = self.upload_to_strapi(blog_data)
                    results.append({
                        'filename': filename,
                        'title': blog_data['title'],
                        'label': blog_data['label'],
                        'success': result is not None,
                        'response': result
                    })
                else:
                    results.append({
                        'filename': filename,
                        'title': 'Failed to parse',
                        'label': '',
                        'success': False,
                        'response': None
                    })
        
        return results
